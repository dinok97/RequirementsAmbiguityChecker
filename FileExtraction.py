import docx
import nltk


class ExtractFile:

    file_directory = ''
    document = ''
    extraction_result = []
    table_header = []
    table_content = []

    def __init__(self, file_directory):
        self.file_directory = file_directory
        self.document = docx.Document(self.file_directory)

        print('\n>>> start extract file')

    def start_extraction(self):
        # table_information contain information about:
        # 1. index table ; 2. requirements column 3. requirements sentence column
        table_information = self.find_table()

        if table_information != "table not found":
            table_index = table_information[0]
            id_req_column = table_information[1]
            req_sentences_column = table_information[2]

            # get table content from selected table
            table_content = self.get_table_content(table_index)
            print('>>> table content :\n%s ' % table_content)

            self.extraction_result.clear()
            for i in range(len(table_content)):
                temp_data = []
                id_req = table_content[i][id_req_column]
                req_sentences = table_content[i][req_sentences_column]
                temp_data.append(id_req)
                temp_data.append(req_sentences)
                self.extraction_result.append(temp_data)
            print('>>> return id and req. sentences :\n%s' % self.extraction_result)

        else:
            print('>>> table not found')
            self.extraction_result.clear()
            self.extraction_result.append("table not found")

    def get_extraction_result(self):
        if self.extraction_result[0] != "table not found":
            return self.extraction_result
        else:
            return "table not found"

    def find_table(self):
        table_information = []
        table_index = 0
        for i in range(len(self.document.tables)):
            header = self.get_table_header(i)
            self.table_header = self.lower_case(header)

            id_req = 0
            for j in range(len(self.table_header)):
                if (self.table_header[j] == 'kode kebutuhan') or \
                        (self.table_header[j] == 'kode kebutuhan sistem') or \
                        (self.table_header[j] == 'kode fungsi') or \
                        (self.table_header[j] == 'kode fungsi sistem') or \
                        (self.table_header[j] == 'id') or \
                        (self.table_header[j] == 'id.') or \
                        (self.table_header[j] == 'uid') or \
                        (self.table_header[j] == 'id kebutuhan') or \
                        (self.table_header[j] == 'id. kebutuhan') or \
                        (self.table_header[j] == 'id kebutuhan sistem') or \
                        (self.table_header[j] == 'id. kebutuhan sistem') or \
                        (self.table_header[j] == 'nomor kebutuhan') or \
                        (self.table_header[j] == 'nomor kebutuhan sistem'):

                    req_sentences = 0
                    for k in range(len(self.table_header)):
                        if (self.table_header[k] == 'kalimat kebutuhan') or \
                                (self.table_header[k] == 'kalimat kebutuhan sistem') or \
                                (self.table_header[k] == 'deskripsi') or \
                                (self.table_header[k] == 'deskripsi fungsi') or \
                                (self.table_header[k] == 'deskripsi kebutuhan') or \
                                (self.table_header[k] == 'deskripsi kebutuhan sistem') or \
                                (self.table_header[k] == 'kebutuhan') or \
                                (self.table_header[k] == 'kebutuhan sistem'):
                            print('>>> find req. table in SKPL')

                            table_information.append(table_index)
                            table_information.append(id_req)
                            table_information.append(req_sentences)

                            return table_information

                        req_sentences += 1
                id_req += 1
            table_index += 1

        print('>>>table not found')
        return "table not found"

    def get_table_content(self, table_index):
        index = 0
        table_content = []
        table = self.document.tables[table_index]

        for rw in table.rows:
            temp_content = []
            for cell in rw.cells:
                temp_content.append(cell.text)
            if index != 0:  # if == 0 its mean header
                table_content.append(temp_content)
            index += 1
        return table_content

    def get_table_header(self, index):
        table_header = []
        table = self.document.tables[index]
        index = 0
        for row in table.rows:
            for cell in row.cells:
                if index == 0:
                    table_header.append(cell.text)
            index += 1
            break
        return table_header

    @staticmethod
    def lower_case(data):
        for i in range(len(data)):
            data[i] = data[i].lower()
        return data

    """def read_documents(self):

        complete_text = []
        for paragraph in self.document.paragraphs:
            # par = paragraph.text.lower().strip()
            if paragraph.text != "":
                complete_text.append(paragraph.text)
        return '\n'.join(complete_text)"""


'''if __name__ == '__main__':
    test = ExtractFile('D:\\PBA.docx')
    # print(test.read_documents())
    # print(test.get_table_content(0))
    # print(test.get_table_header(0))
    print(test.get_requirements_data())
    # print(tableHeader)'''