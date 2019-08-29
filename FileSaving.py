from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION


class ExportFile:

    srs_file = ''
    document = ''
    export_file_name = ''
    srs_source_name = ''
    detection_time = ''

    clear_requirements = 0
    ambiguous_requirements = 0
    total_requirements = 0

    detection_result = []

    def __init__(self, export_file_name, srs_file):

        # initialize the export file name
        self.export_file_name = export_file_name

        # initialize srs object
        self.srs_file = srs_file

        # formatting text to display
        self.srs_source_name = (': %s' % self.srs_file.get_file_name())
        self.detection_time = (': %s' % self.srs_file.get_correction_time())
        self.clear_requirements = (': %s Kebutuhan' % self.srs_file.get_normal_req())
        self.ambiguous_requirements = (': %s Kebutuhan' % self.srs_file.get_ambiguous_req())
        self.total_requirements = (': %s Kebutuhan' % self.srs_file.get_total_req())
        self.detection_result = self.srs_file.get_detection_result()

        # show detail information srs document and detection result
        print('\n>>> user clicked export file')
        print('direktori dokumen %s' % self.srs_source_name)
        print('waktu deteksi %s' % self.detection_time)
        print('kabutuhan normal %s' % self.clear_requirements)
        print('kebutuhan ambigu %s' % self.ambiguous_requirements)
        print('kebutuhan total %s' % self.total_requirements)
        print('hasil deteksi : %s' % self.detection_result)

        # initialize document from Document class in docx library
        self.document = Document()

    def start_export_file(self):
        # setting size and orientation of ms. word paper
        section = self.document.sections[0]
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
        print('>>> paper orientation : %s' % section.orientation)

        # add heading in document
        self.document.add_heading('Detail Hasil Pemeriksaan Kalimat Kebutuhan', 0)

        # array that contains information about srs file already detected
        doc_information = [
            ['Direktori Dokumen', self.srs_source_name],
            ['Waktu Deteksi', self.detection_time],
            ['Kebutuhan Sesuai', self.clear_requirements],
            ['Kebutuhan Ambigu', self.ambiguous_requirements],
            ['Total Kebutuhan', self.total_requirements]
        ]

        # information_table contain information about srs document
        information_table = self.document.add_table(rows=0, cols=2)
        for name, content in doc_information:
            # format column one
            row_cells = information_table.add_row().cells
            row_cells[0].text = name
            row_cells[0].width = Inches(2.0)
            paragraph = row_cells[0].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.runs
            font = run[0].font
            font.name = 'Calibri'
            font.size = Pt(12)

            # format column two
            row_cells[1].text = content
            row_cells[1].width = Inches(7)
            paragraph = row_cells[1].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            run = paragraph.runs
            run[0].bold = True
            font = run[0].font
            font.name = 'Calibri'
            font.size = Pt(12)

        # text to show table title
        table_title = self.document.add_paragraph('')
        run = table_title.add_run('Tabel Hasil Deteksi Kalimat Kebutuhan')
        run.bold = True
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(15)
        table_title.paragraph_format.space_before = Pt(15)
        table_title.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # table to show detection result
        table = self.document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No.'
        hdr_cells[0].width = Inches(0.4)

        hdr_cells[1].text = 'Id Keb.'
        hdr_cells[1].width = Inches(0.7)

        hdr_cells[2].text = 'Kalimat Kebutuhan'
        hdr_cells[2].width = Inches(5.0)

        hdr_cells[3].text = 'Ambiguitas'
        hdr_cells[3].width = Inches(0.5)

        hdr_cells[4].text = 'Pola Ambigu'
        hdr_cells[4].width = Inches(1.0)

        hdr_cells[5].text = 'Rekomendasi'
        hdr_cells[5].width = Inches(2.0)

        # this loop used to insert detection result data into table
        for no, id_rec, rec, abg, loc_abg, sgstn in self.detection_result:
            row_cells = table.add_row().cells
            row_cells[0].text = str(no)
            row_cells[0].width = Inches(0.4)
            row_cells[1].text = id_rec
            row_cells[1].width = Inches(0.7)
            row_cells[2].text = rec
            row_cells[2].width = Inches(5.0)
            row_cells[3].text = abg
            row_cells[3].width = Inches(0.5)
            row_cells[4].text = loc_abg
            row_cells[4].width = Inches(1.0)
            row_cells[5].text = sgstn
            row_cells[5].width /= Inches(2.0)

        # method to set style to the table
        self.add_table_style(table)

    def save_file(self):
        self.document.save('%s.docx' % self.export_file_name)

        # print successful message
        print(">>> successful to save file at %s.docx" % self.export_file_name)

    @staticmethod
    def add_table_style(table):
        table.style = 'TableGrid'
        table.autofit = False
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.name = 'Calibri'


"""docIdentity = document.add_paragraph('')
        docIdentity.add_run('Nama Dokumen         :').font.name = 'Calibri'
        run = docIdentity.add_run('SKPL2.docx')
        run.bold = True
        font = run.font
        font.name = 'Calibri'

        docIdentity.add_run('\nWaktu Deteksi      :').font.name = 'Calibri'
        run = docIdentity.add_run('Kamis, 15 Mei 2019')
        run.bold = True
        fonts = run.font
        fonts.name = 'Calibri'

        docIdentity.add_run('\nKebutuhan Sesuai   :').font.name = 'Calibri'
        run = docIdentity.add_run('15')
        run.bold = True
        fonts = run.font
        fonts.name = 'Calibri'

        docIdentity.add_run('\nKebutuhan Ambigu   :').font.name = 'Calibri'
        run = docIdentity.add_run('12')
        run.bold = True
        fonts = run.font
        fonts.name = 'Calibri'

        docIdentity.add_run('\nTotal Kebutuhan    :').font.name = 'Calibri'
        run = docIdentity.add_run('27')
        run.bold = True
        fonts = run.font
        fonts.name = 'Calibri'"""